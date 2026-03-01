import io
from typing import List
from core.models import DocumentChunk, FileType
from .base_extractor import BaseExtractor

class WordExtractor(BaseExtractor):
    SUPPORTED_EXTENSIONS = ["docx", "doc"]

    def extract(self, file_path: str) -> List[DocumentChunk]:
        chunks: List[DocumentChunk] = []
        source = file_path.split("/")[-1]
        chunk_idx = 0

        try:
            from docx import Document
            from docx.oxml.ns import qn
            import zipfile

            doc = Document(file_path)

            page_size = 50
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            for page_num, start in enumerate(range(0, max(len(paragraphs), 1), page_size), start=1):
                page_text = "\n".join(paragraphs[start: start + page_size])
                if page_text:
                    chunks.append(self._make_text_chunk(
                        text=f"[Section {page_num}]\n{page_text}",
                        source=source,
                        page=page_num,
                        chunk_index=chunk_idx,
                        file_type=FileType.WORD,
                    ))
                    chunk_idx += 1

            for tbl_idx, table in enumerate(doc.tables):
                rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                table_text = "\n".join(rows)
                if table_text.strip():
                    chunks.append(self._make_text_chunk(
                        text=f"[Table {tbl_idx + 1}]\n{table_text}",
                        source=source,
                        page=0,
                        chunk_index=chunk_idx,
                        file_type=FileType.WORD,
                    ))
                    chunk_idx += 1

            with zipfile.ZipFile(file_path, "r") as zf:
                image_files = [n for n in zf.namelist() if n.startswith("word/media/")]
                for img_path in image_files:
                    ext = img_path.rsplit(".", 1)[-1].lower()
                    mime = f"image/{ext}" if ext != "jpg" else "image/jpeg"
                    image_data = zf.read(img_path)
                    chunks.append(self._make_image_chunk(
                        image_data=image_data,
                        mime=mime,
                        source=source,
                        page=0,
                        chunk_index=chunk_idx,
                        file_type=FileType.WORD,
                    ))
                    chunk_idx += 1

        except ImportError as e:
            raise RuntimeError(
                f"Missing dependency for Word extraction: {e}. "
                "Run: pip install python-docx"
            )
        except Exception as e:
            raise RuntimeError(f"Word extraction failed for '{file_path}': {e}")

        return chunks 